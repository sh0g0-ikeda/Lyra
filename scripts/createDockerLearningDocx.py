from __future__ import annotations

from pathlib import Path
import re
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "assets" / "docker-learning"
OUT_DOCX = ROOT / "docs" / "docker-learning-lyra.docx"


def font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/meiryo.ttc",
        "C:/Windows/Fonts/msgothic.ttc",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], width: int, fnt, fill=(20, 20, 20), line_gap: int = 8) -> int:
    x, y = xy
    current = ""
    lines: list[str] = []
    for char in text:
        test = current + char
        bbox = draw.textbbox((0, 0), test, font=fnt)
        if bbox[2] - bbox[0] > width and current:
            lines.append(current)
            current = char
        else:
            current = test
    if current:
        lines.append(current)
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + line_gap
    return y


def box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, fnt, width: int = 2) -> None:
    draw.rounded_rectangle(xy, radius=16, outline=(40, 40, 40), width=width, fill=(255, 255, 255))
    x1, y1, x2, y2 = xy
    y = y1 + 18
    draw_wrapped(draw, text, (x1 + 18, y), x2 - x1 - 36, fnt)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill=(35, 35, 35), width=3)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - direction * 14, ey - 8), (ex - direction * 14, ey + 8)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 8, ey - direction * 14), (ex + 8, ey - direction * 14)]
    draw.polygon(points, fill=(35, 35, 35))


def create_diagram_lifecycle(path: Path) -> None:
    img = Image.new("RGB", (1500, 520), "white")
    d = ImageDraw.Draw(img)
    title = font(34)
    body = font(28)
    d.text((40, 30), "Dockerの基本の流れ", font=title, fill=(0, 0, 0))
    boxes = [
        (70, 160, 330, 300, "Dockerfile\nイメージの作り方"),
        (450, 160, 710, 300, "イメージ\n実行環境の完成品"),
        (830, 160, 1090, 300, "コンテナ\n実行中の実体"),
        (1210, 160, 1450, 300, "ログ・DB・外部API\n実行結果"),
    ]
    for b in boxes:
        box(d, b[:4], b[4], body)
    arrow(d, (330, 230), (450, 230))
    arrow(d, (710, 230), (830, 230))
    arrow(d, (1090, 230), (1210, 230))
    d.text((355, 195), "build", font=font(24), fill=(0, 0, 0))
    d.text((735, 195), "run", font=font(24), fill=(0, 0, 0))
    d.text((1115, 195), "処理", font=font(24), fill=(0, 0, 0))
    img.save(path)


def create_diagram_local(path: Path) -> None:
    img = Image.new("RGB", (1500, 680), "white")
    d = ImageDraw.Draw(img)
    title = font(34)
    body = font(26)
    d.text((40, 30), "Lyraのローカル開発構成", font=title, fill=(0, 0, 0))
    box(d, (70, 170, 370, 300), "ブラウザ\nhttp://127.0.0.1:5173", body)
    box(d, (500, 120, 850, 250), "Vite dev server\nbun run web:dev\nPC上で起動", body)
    box(d, (500, 360, 850, 500), "Lyra API\nbun run dev\nPC上で起動", body)
    box(d, (1000, 360, 1390, 500), "PostgreSQL\nDocker Compose\nlyra-postgres", body)
    arrow(d, (370, 235), (500, 185))
    arrow(d, (675, 250), (675, 360))
    arrow(d, (850, 430), (1000, 430))
    d.text((885, 390), "DATABASE_URL\n127.0.0.1:5432", font=font(22), fill=(0, 0, 0))
    img.save(path)


def create_diagram_production(path: Path) -> None:
    img = Image.new("RGB", (1500, 820), "white")
    d = ImageDraw.Draw(img)
    title = font(34)
    body = font(24)
    d.text((40, 30), "Lyraの本番Docker構成", font=title, fill=(0, 0, 0))
    box(d, (70, 130, 340, 250), "ブラウザ", body)
    box(d, (450, 130, 720, 250), "CloudFront / ALB", body)
    box(d, (830, 130, 1160, 250), "ECS Fargate\nAPIコンテナ", body)
    box(d, (830, 430, 1160, 560), "ECS Fargate\nWorkerコンテナ", body)
    box(d, (1240, 90, 1450, 200), "RDS\nPostgreSQL", body)
    box(d, (1240, 245, 1450, 355), "S3\n画像保存", body)
    box(d, (500, 430, 720, 560), "SQS\n生成ジョブ", body)
    box(d, (1240, 430, 1450, 560), "OpenAI\n画像生成", body)
    box(d, (1240, 600, 1450, 720), "Secrets Manager\n機密情報", body)
    arrow(d, (340, 190), (450, 190))
    arrow(d, (720, 190), (830, 190))
    arrow(d, (1160, 180), (1240, 145))
    arrow(d, (1160, 210), (1240, 300))
    arrow(d, (970, 250), (610, 430))
    arrow(d, (720, 500), (830, 500))
    arrow(d, (1160, 500), (1240, 500))
    arrow(d, (1000, 560), (1240, 660))
    img.save(path)


def create_diagram_ports(path: Path) -> None:
    img = Image.new("RGB", (1500, 520), "white")
    d = ImageDraw.Draw(img)
    title = font(34)
    body = font(28)
    d.text((40, 30), "ポート割り当ての考え方", font=title, fill=(0, 0, 0))
    box(d, (100, 160, 520, 340), "ホストPC\nlocalhost:5432\nユーザーやAPIが見る入口", body)
    box(d, (940, 160, 1360, 340), "PostgreSQLコンテナ\ncontainer:5432\nDB本体が待ち受ける場所", body)
    arrow(d, (520, 250), (940, 250))
    d.text((625, 205), '"5432:5432"', font=font(30), fill=(0, 0, 0))
    d.text((610, 260), "左がホスト側、右がコンテナ側", font=font(24), fill=(0, 0, 0))
    img.save(path)


def create_images() -> dict[str, Path]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    paths = {
        "lifecycle": OUT_DIR / "docker-lifecycle.png",
        "local": OUT_DIR / "lyra-local-docker.png",
        "production": OUT_DIR / "lyra-production-docker.png",
        "ports": OUT_DIR / "docker-port-mapping.png",
    }
    create_diagram_lifecycle(paths["lifecycle"])
    create_diagram_local(paths["local"])
    create_diagram_production(paths["production"])
    create_diagram_ports(paths["ports"])
    return paths


def set_document_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Yu Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(10.5)
    normal.font.bold = False

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Yu Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.bold = False
    doc.styles["Title"].font.size = Pt(20)
    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 3"].font.size = Pt(12)


def add_paragraph(doc: Document, text: str = "", style: str | None = None, alignment=None) -> None:
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.bold = False
    run.font.name = "Yu Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")


def add_code(doc: Document, code: str) -> None:
    for line in code.rstrip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.bold = False


def add_table_with_caption(doc: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    add_paragraph(doc, caption, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = False
                    run.font.name = "Yu Gothic"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")


def add_figure(doc: Document, image_path: Path, caption: str, width_cm: float = 15.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    add_paragraph(doc, caption, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def add_math(doc: Document, expression: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omath_para = OxmlElement("m:oMathPara")
    omath = OxmlElement("m:oMath")
    mr = OxmlElement("m:r")
    mt = OxmlElement("m:t")
    mt.text = expression
    mr.append(mt)
    omath.append(mr)
    omath_para.append(omath)
    p._p.append(omath_para)


def clean_inline(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    return text


def parse_markdown_subset(doc: Document, md_path: Path, image_paths: dict[str, Path]) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    table_no = 1
    figure_no = 1
    pending_table: list[str] = []
    in_code = False
    code_lines: list[str] = []
    list_counter = 1

    def flush_table() -> None:
        nonlocal pending_table, table_no
        if not pending_table:
            return
        rows = []
        for row in pending_table:
            cells = [clean_inline(cell.strip()) for cell in row.strip().strip("|").split("|")]
            rows.append(cells)
        pending_table = []
        if len(rows) < 2:
            return
        headers = rows[0]
        data_rows = [row for row in rows[2:] if len(row) == len(headers)]
        add_table_with_caption(doc, f"表{table_no} Docker学習項目", headers, data_rows)
        table_no += 1

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            add_code(doc, "\n".join(code_lines))
            code_lines = []

    for line in lines:
        if line.strip().startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                flush_table()
                in_code = True
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            pending_table.append(line)
            continue
        flush_table()

        raw = line.rstrip()
        if not raw or raw == "---":
            list_counter = 1
            continue

        if raw.startswith("# "):
            add_paragraph(doc, clean_inline(raw[2:].strip()), "Title", WD_ALIGN_PARAGRAPH.CENTER)
            add_paragraph(doc, "Docker公式Get StartedとLyraの実装を対応させて理解するための資料", alignment=WD_ALIGN_PARAGRAPH.CENTER)
            doc.add_page_break()
            continue
        if raw.startswith("## "):
            title = clean_inline(raw[3:].strip())
            add_paragraph(doc, title, "Heading 1")
            if title.startswith("3. Dockerの基本概念"):
                add_figure(doc, image_paths["lifecycle"], f"図{figure_no} Dockerfile、イメージ、コンテナの関係")
                figure_no += 1
            elif title.startswith("6. LyraのDocker構成全体"):
                add_figure(doc, image_paths["local"], f"図{figure_no} Lyraのローカル開発構成")
                figure_no += 1
                add_figure(doc, image_paths["production"], f"図{figure_no} Lyraの本番Docker構成")
                figure_no += 1
            elif title.startswith("27. Dockerとポート"):
                add_figure(doc, image_paths["ports"], f"図{figure_no} ホスト側ポートとコンテナ側ポートの対応")
                figure_no += 1
            continue
        if raw.startswith("### "):
            add_paragraph(doc, clean_inline(raw[4:].strip()), "Heading 2")
            continue

        stripped = raw.lstrip()
        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            add_paragraph(doc, f"{list_counter}. {clean_inline(bullet_match.group(1).strip())}")
            list_counter += 1
            continue
        numbered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered_match:
            add_paragraph(doc, clean_inline(stripped))
            continue

        add_paragraph(doc, clean_inline(raw))

    flush_table()
    flush_code()


def add_equation_section(doc: Document) -> None:
    add_paragraph(doc, "40. Word上で数式として確認できるDockerの基本式", "Heading 1")
    add_paragraph(doc, "この章の式は、通常の文字列ではなくWordの数式オブジェクトとして挿入しています。Dockerの理解に必要な関係を、最小限の式として示します。")
    add_paragraph(doc, "ホスト側ポートとコンテナ側ポートの対応は次の形で表せます。")
    add_math(doc, "host_port : container_port = 5432 : 5432")
    add_paragraph(doc, "Dockerfile、イメージ、コンテナの関係は次のように表せます。")
    add_math(doc, "Dockerfile → Image → Container")
    add_paragraph(doc, "Docker buildの再現性は、依存定義とlockfileが一致していることに依存します。")
    add_math(doc, "Reproducible Build = package.json + bun.lock + frozen lockfile")


def main() -> None:
    md_path = ROOT / "docs" / "docker-learning-lyra.md"
    if not md_path.exists():
        raise FileNotFoundError(md_path)

    image_paths = create_images()
    doc = Document()
    set_document_style(doc)

    parse_markdown_subset(doc, md_path, image_paths)
    add_equation_section(doc)

    add_paragraph(doc, "付録. 生成条件の確認", "Heading 1")
    add_paragraph(doc, "この文書では、表の表題は表の上、図の表題は図の下に配置しています。箇条書き記号は使わず、番号付きの通常段落として整理しています。装飾目的の色、網掛け、太字は使用していません。")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
